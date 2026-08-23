# -*- coding: utf-8 -*-
"""
=============================================================================
 mise_en_page.py
-----------------------------------------------------------------------------
 Briques de mise en forme du rapport AgriVision-AI.

 Ce fichier ne contient AUCUN texte du rapport. Il ne sait que fabriquer des
 objets Word propres : titres, paragraphes, tableaux, encadres, figures.
 Le texte, lui, vit dans rapport/chapitres/*.md, en Markdown, pour que les
 deux membres du binome puissent rediger sans ecrire une ligne de Python.

 Copie du generateur du guide de projet (build/mise_en_page.py), amenee dans
 le depot pour que le rapport soit reproductible par n'importe qui a partir
 d'un simple clone. Seuls les chemins de figures changent : ils pointent
 desormais vers reports/ et docs/.

 Document en noir et blanc, sobre, sans effet decoratif. Seuls deux gris de
 fond subsistent, pour les en-tetes de tableau et les encadres.
=============================================================================
"""

from pathlib import Path

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# =============================================================================
#  Reperes typographiques (aucune couleur)
# =============================================================================

NOIR = RGBColor(0x00, 0x00, 0x00)
TEXTE_COUL = RGBColor(0x1A, 0x1A, 0x1A)      # noir legerement adouci, plus lisible
GRIS_DOUX = RGBColor(0x59, 0x59, 0x59)       # legendes, pied de page

FOND_ENTETE = "D9D9D9"      # en-tete de tableau
FOND_LIGNE = "F4F4F4"       # une ligne sur deux
FOND_ENCADRE = "F2F2F2"     # encadre de vigilance
FILET = "808080"            # gris moyen pour tous les filets

TEXTE = "Calibri"
MONO = "Consolas"

LARGEUR_UTILE_CM = 17.0     # 21 cm moins deux marges de 2 cm

RACINE = Path(__file__).resolve().parents[2]         # .../repo
DOSSIER_FIGURES = RACINE / "reports"                 # les figures du rapport
DOSSIER_DOCS = RACINE / "docs"                       # schema d'architecture


# =============================================================================
#  Compteurs partages (figures et tableaux)
# =============================================================================

class Compteurs:
    """Numerotation continue des figures et des tableaux sur tout le document."""

    def __init__(self):
        self.figure = 0
        self.tableau = 0
        self.index_figures = []
        self.index_tableaux = []

    def figure_suivante(self, legende_courte):
        self.figure += 1
        self.index_figures.append((self.figure, legende_courte))
        return self.figure

    def tableau_suivant(self, legende_courte):
        self.tableau += 1
        self.index_tableaux.append((self.tableau, legende_courte))
        return self.tableau


CPT = Compteurs()


# =============================================================================
#  Descente dans le XML de Word (python-docx ne sait pas tout faire seul)
# =============================================================================

def _fond(paragraphe, couleur_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), couleur_hex)
    paragraphe._p.get_or_add_pPr().append(shd)


def _fond_cellule(cellule, couleur_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), couleur_hex)
    cellule._tc.get_or_add_tcPr().append(shd)


def _bordure(paragraphe, cotes, couleur_hex, epaisseur=6, espace=6):
    pPr = paragraphe._p.get_or_add_pPr()
    bordures = pPr.find(qn("w:pBdr"))
    if bordures is None:
        bordures = OxmlElement("w:pBdr")
        pPr.append(bordures)
    for cote in cotes:
        b = OxmlElement(f"w:{cote}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(epaisseur))
        b.set(qn("w:space"), str(espace))
        b.set(qn("w:color"), couleur_hex)
        bordures.append(b)


def _sans_coupure(paragraphe, avec_suivant=True):
    """Empeche Word de couper ce paragraphe, et de le separer du suivant."""
    pPr = paragraphe._p.get_or_add_pPr()
    balises = ["w:keepLines"] + (["w:keepNext"] if avec_suivant else [])
    for balise in balises:
        pPr.append(OxmlElement(balise))


def _champ(paragraphe, instruction, taille=8, couleur=GRIS_DOUX, gras=False):
    """Insere un champ Word calcule (PAGE, NUMPAGES, TOC)."""
    r = paragraphe.add_run()
    r.font.size = Pt(taille)
    r.font.color.rgb = couleur
    r.font.bold = gras
    r.font.name = TEXTE

    debut = OxmlElement("w:fldChar")
    debut.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separateur = OxmlElement("w:fldChar")
    separateur.set(qn("w:fldCharType"), "separate")
    fin = OxmlElement("w:fldChar")
    fin.set(qn("w:fldCharType"), "end")

    r._r.append(debut)
    r._r.append(instr)
    r._r.append(separateur)
    r._r.append(fin)
    return r


# =============================================================================
#  Paragraphes
# =============================================================================

def para(doc, texte="", taille=10.5, gras=False, italique=False,
         couleur=TEXTE_COUL, avant=2, apres=6, police=TEXTE,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, retrait=0, interligne=1.08):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(avant)
    pf.space_after = Pt(apres)
    pf.line_spacing = interligne
    if retrait:
        pf.left_indent = Cm(retrait)
    if align is not None:
        p.alignment = align
    if texte:
        r = p.add_run(texte)
        r.font.size = Pt(taille)
        r.font.bold = gras
        r.font.italic = italique
        r.font.color.rgb = couleur
        r.font.name = police
    return p


def riche(doc, morceaux, taille=10.5, avant=2, apres=6, retrait=0,
          align=WD_ALIGN_PARAGRAPH.JUSTIFY, interligne=1.08):
    """
    Paragraphe compose de plusieurs morceaux :
        [("texte normal", {}), ("en gras", {"gras": True})]
    """
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(avant)
    pf.space_after = Pt(apres)
    pf.line_spacing = interligne
    if retrait:
        pf.left_indent = Cm(retrait)
    if align is not None:
        p.alignment = align
    for texte, opts in morceaux:
        r = p.add_run(texte)
        r.font.size = Pt(opts.get("taille", taille))
        r.font.bold = opts.get("gras", False)
        r.font.italic = opts.get("italique", False)
        r.font.color.rgb = opts.get("couleur", TEXTE_COUL)
        r.font.name = opts.get("police", TEXTE)
    return p


def puce(doc, texte, gras_debut=None, retrait=0.7, taille=10.5):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(4)
    pf.left_indent = Cm(retrait)
    pf.first_line_indent = Cm(-0.35)
    pf.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    r = p.add_run("•\t")
    r.font.size = Pt(taille)
    r.font.color.rgb = TEXTE_COUL
    r.font.name = TEXTE
    if gras_debut:
        rg = p.add_run(gras_debut)
        rg.font.size = Pt(taille)
        rg.font.bold = True
        rg.font.color.rgb = TEXTE_COUL
        rg.font.name = TEXTE
    r2 = p.add_run(texte)
    r2.font.size = Pt(taille)
    r2.font.color.rgb = TEXTE_COUL
    r2.font.name = TEXTE
    return p


def numero(doc, indice, texte, gras_debut=None, retrait=0.9, taille=10.5):
    """Element d'une liste numerotee ecrite a la main (numerotation stable)."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(4)
    pf.left_indent = Cm(retrait)
    pf.first_line_indent = Cm(-0.9)
    pf.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    r = p.add_run(f"{indice}.\t")
    r.font.size = Pt(taille)
    r.font.bold = True
    r.font.color.rgb = TEXTE_COUL
    r.font.name = TEXTE
    if gras_debut:
        rg = p.add_run(gras_debut)
        rg.font.size = Pt(taille)
        rg.font.bold = True
        rg.font.color.rgb = TEXTE_COUL
        rg.font.name = TEXTE
    r2 = p.add_run(texte)
    r2.font.size = Pt(taille)
    r2.font.color.rgb = TEXTE_COUL
    r2.font.name = TEXTE
    return p


# =============================================================================
#  Titres (styles Word natifs, indispensables au sommaire automatique)
# =============================================================================

def _titre(doc, texte, niveau, taille, avant, apres, majuscules=False):
    p = doc.add_paragraph(style=f"Heading {niveau}")
    pf = p.paragraph_format
    pf.space_before = Pt(avant)
    pf.space_after = Pt(apres)
    pf.line_spacing = 1.05
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(texte.upper() if majuscules else texte)
    r.font.size = Pt(taille)
    r.font.bold = True
    r.font.color.rgb = NOIR
    r.font.name = TEXTE
    _sans_coupure(p)
    return p


def _filet_sous_titre(doc, apres=12, epaisseur=8):
    trait = doc.add_paragraph()
    trait.paragraph_format.space_before = Pt(0)
    trait.paragraph_format.space_after = Pt(apres)
    _bordure(trait, ["bottom"], FILET, epaisseur=epaisseur, espace=1)
    _sans_coupure(trait)
    return trait


def chapitre(doc, numero_chapitre, texte, saut=True):
    """Ouverture d'un chapitre : saut de page, mention du numero, titre, filet."""
    if saut:
        doc.add_page_break()

    bandeau = doc.add_paragraph()
    bandeau.paragraph_format.space_before = Pt(0)
    bandeau.paragraph_format.space_after = Pt(2)
    r = bandeau.add_run(f"CHAPITRE {numero_chapitre}")
    r.font.size = Pt(9.5)
    r.font.bold = True
    r.font.color.rgb = GRIS_DOUX
    r.font.name = TEXTE
    _sans_coupure(bandeau)

    p = _titre(doc, texte, 1, 17, avant=0, apres=4)
    _filet_sous_titre(doc)
    return p


def titre1(doc, texte, saut=True):
    """Titre de niveau chapitre sans numerotation (glossaire, sommaire)."""
    if saut:
        doc.add_page_break()
    p = _titre(doc, texte, 1, 19, avant=0, apres=4)
    _filet_sous_titre(doc)
    return p


def titre2(doc, texte):
    return _titre(doc, texte, 2, 12.5, avant=12, apres=4)


def titre3(doc, texte):
    return _titre(doc, texte, 3, 11, avant=10, apres=3)


def intertitre(doc, texte):
    """Petit intertitre non repris dans le sommaire."""
    p = para(doc, texte, taille=10, gras=True, couleur=GRIS_DOUX,
             avant=10, apres=2, align=WD_ALIGN_PARAGRAPH.LEFT)
    _sans_coupure(p)
    return p


# =============================================================================
#  Encadre
# =============================================================================

def encadre(doc, etiquette, texte, taille=10):
    """
    Bloc mis en retrait : fond gris tres clair, filet gris sur les quatre cotes.
    L'etiquette est facultative et s'affiche en gras au debut du texte.
    """
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(10)
    pf.left_indent = Cm(0.2)
    pf.right_indent = Cm(0.2)
    pf.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if etiquette:
        r = p.add_run(f"{etiquette}   ")
        r.font.size = Pt(taille - 0.5)
        r.font.bold = True
        r.font.color.rgb = NOIR
        r.font.name = TEXTE
    r = p.add_run(texte)
    r.font.size = Pt(taille)
    r.font.color.rgb = TEXTE_COUL
    r.font.name = TEXTE
    _fond(p, FOND_ENCADRE)
    _bordure(p, ["top", "left", "bottom", "right"], FILET, epaisseur=6, espace=8)
    _sans_coupure(p, avec_suivant=False)
    return p


def a_retenir(doc, texte):
    return encadre(doc, "À RETENIR", texte)


def vigilance(doc, texte):
    return encadre(doc, "ATTENTION", texte)


# =============================================================================
#  Tableaux
# =============================================================================

def tableau(doc, entetes, lignes, largeurs=None, legende=None, taille=9):
    """
    Tableau a filets fins, en-tete en gras sur fond gris clair.
    Une valeur entouree de barres obliques inverses passe en chasse fixe :
        "`classes.json`"
    Une valeur entouree de deux asterisques passe en gras.
    """
    t = doc.add_table(rows=1, cols=len(entetes))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    if largeurs:
        t.autofit = False
        layout = OxmlElement("w:tblLayout")
        layout.set(qn("w:type"), "fixed")
        t._tbl.tblPr.append(layout)

    # L'en-tete se repete en haut de chaque page si le tableau est coupe.
    trPr = t.rows[0]._tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:tblHeader"))

    for i, texte in enumerate(entetes):
        cell = t.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(texte)
        r.font.size = Pt(taille)
        r.font.bold = True
        r.font.color.rgb = NOIR
        r.font.name = TEXTE
        _fond_cellule(cell, FOND_ENTETE)

    for idx, ligne in enumerate(lignes):
        cells = t.add_row().cells
        for i, valeur in enumerate(ligne):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.08
            valeur = str(valeur)
            fixe = valeur.startswith("`") and valeur.endswith("`") and len(valeur) > 1
            contenu = valeur.strip("`")
            gras = contenu.startswith("**") and contenu.endswith("**")
            if gras:
                contenu = contenu.strip("*")
            r = p.add_run(contenu)
            r.font.size = Pt(taille - 0.5 if fixe else taille)
            r.font.name = MONO if fixe else TEXTE
            r.font.bold = gras
            r.font.color.rgb = TEXTE_COUL
            if idx % 2 == 1:
                _fond_cellule(cells[i], FOND_LIGNE)

    # Aucune ligne ne doit etre coupee entre deux pages.
    for r_ in t.rows:
        r_._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))

    # Un tableau court reste solidaire du texte qui le precede.
    if len(lignes) <= 3:
        for r_ in t.rows[:-1]:
            for cellule in r_.cells:
                for p_ in cellule.paragraphs:
                    _sans_coupure(p_)

    if largeurs:
        for r_ in t.rows:
            for i, l in enumerate(largeurs):
                r_.cells[i].width = Cm(l)
        for i, l in enumerate(largeurs):
            t.columns[i].width = Cm(l)

    if legende:
        num = CPT.tableau_suivant(legende)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(f"Tableau {num}. ")
        r.font.size = Pt(9)
        r.font.bold = True
        r.font.color.rgb = NOIR
        r.font.name = TEXTE
        r2 = p.add_run(legende)
        r2.font.size = Pt(9)
        r2.font.italic = True
        r2.font.color.rgb = GRIS_DOUX
        r2.font.name = TEXTE
    else:
        para(doc, "", taille=4, avant=0, apres=5)
    return t


# =============================================================================
#  Figures
# =============================================================================

def figure(doc, nom_fichier, legende, largeur_cm=15.0, dossier=None, cadre=False):
    """Insere une image et sa legende numerotee. Renvoie le numero attribue."""
    dossier = dossier or DOSSIER_FIGURES
    chemin = Path(dossier) / nom_fichier
    if not chemin.is_file():
        raise FileNotFoundError(f"Image introuvable : {chemin}")

    num = CPT.figure_suivante(legende)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(chemin), width=Cm(largeur_cm))
    if cadre:
        _bordure(p, ["top", "left", "bottom", "right"], FILET, epaisseur=4, espace=2)
    _sans_coupure(p)

    pl = doc.add_paragraph()
    pl.paragraph_format.space_before = Pt(4)
    pl.paragraph_format.space_after = Pt(12)
    pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pl.paragraph_format.line_spacing = 1.05
    r1 = pl.add_run(f"Figure {num}. ")
    r1.font.size = Pt(9)
    r1.font.bold = True
    r1.font.color.rgb = NOIR
    r1.font.name = TEXTE
    r2 = pl.add_run(legende)
    r2.font.size = Pt(9)
    r2.font.italic = True
    r2.font.color.rgb = GRIS_DOUX
    r2.font.name = TEXTE
    return num


# =============================================================================
#  Document, page de garde, sommaire
# =============================================================================

def _regler_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = TEXTE
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TEXTE_COUL
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), TEXTE)

    for niveau, taille in ((1, 19), (2, 13.5), (3, 11.5)):
        style = doc.styles[f"Heading {niveau}"]
        style.font.name = TEXTE
        style.font.size = Pt(taille)
        style.font.bold = True
        style.font.color.rgb = NOIR
        style.element.rPr.rFonts.set(qn("w:eastAsia"), TEXTE)


def _mettre_a_jour_les_champs(doc):
    """
    Demande a Word de recalculer tous les champs a l'ouverture du fichier.

    Sans cela, le sommaire reste vide tant que personne n'a fait Ctrl+A puis F9.
    Comme le document est destine a etre envoye a quelqu'un d'autre, on ne peut
    pas compter sur cette manipulation.
    """
    parametres = doc.settings.element
    balise = OxmlElement("w:updateFields")
    balise.set(qn("w:val"), "true")
    parametres.append(balise)


def nouveau_document(titre_pied):
    doc = __import__("docx").Document()
    _mettre_a_jour_les_champs(doc)
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)

    _regler_styles(doc)

    pied = section.footer.paragraphs[0]
    pied.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pied.add_run(f"{titre_pied}    |    page ")
    r.font.size = Pt(8)
    r.font.color.rgb = GRIS_DOUX
    r.font.name = TEXTE
    _champ(pied, "PAGE")
    r2 = pied.add_run(" sur ")
    r2.font.size = Pt(8)
    r2.font.color.rgb = GRIS_DOUX
    r2.font.name = TEXTE
    _champ(pied, "NUMPAGES")
    return doc


def page_de_garde(doc, etablissement, filiere, matiere, nature, titre,
                  sous_titre, auteurs, encadrant, date_rendu):
    """Page de garde sobre : filets gris, titre noir, auteurs, encadrant, date."""
    para(doc, "", taille=8, avant=0, apres=14)

    para(doc, etablissement, taille=12, gras=True, couleur=NOIR,
         align=WD_ALIGN_PARAGRAPH.CENTER, avant=0, apres=2)
    para(doc, filiere, taille=11, couleur=GRIS_DOUX,
         align=WD_ALIGN_PARAGRAPH.CENTER, avant=0, apres=2)
    para(doc, matiere, taille=11, italique=True, couleur=GRIS_DOUX,
         align=WD_ALIGN_PARAGRAPH.CENTER, avant=0, apres=30)

    trait = doc.add_paragraph()
    trait.paragraph_format.space_before = Pt(0)
    trait.paragraph_format.space_after = Pt(18)
    _bordure(trait, ["bottom"], FILET, epaisseur=12, espace=1)

    para(doc, nature, taille=11, gras=True, couleur=GRIS_DOUX,
         align=WD_ALIGN_PARAGRAPH.CENTER, avant=0, apres=10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(titre)
    r.font.size = Pt(26)
    r.font.bold = True
    r.font.color.rgb = NOIR
    r.font.name = TEXTE

    para(doc, sous_titre, taille=12.5, italique=True, couleur=GRIS_DOUX,
         align=WD_ALIGN_PARAGRAPH.CENTER, avant=0, apres=18)

    trait2 = doc.add_paragraph()
    trait2.paragraph_format.space_before = Pt(0)
    trait2.paragraph_format.space_after = Pt(26)
    _bordure(trait2, ["bottom"], FILET, epaisseur=12, espace=1)

    para(doc, "Rédigé par", taille=10, couleur=GRIS_DOUX,
         align=WD_ALIGN_PARAGRAPH.CENTER, avant=10, apres=4)
    for auteur in auteurs:
        para(doc, auteur, taille=13, gras=True, couleur=TEXTE_COUL,
             align=WD_ALIGN_PARAGRAPH.CENTER, avant=0, apres=3)

    para(doc, "Sous la direction de", taille=10, couleur=GRIS_DOUX,
         align=WD_ALIGN_PARAGRAPH.CENTER, avant=22, apres=4)
    para(doc, encadrant, taille=13, gras=True, couleur=TEXTE_COUL,
         align=WD_ALIGN_PARAGRAPH.CENTER, avant=0, apres=28)

    para(doc, date_rendu, taille=10.5, couleur=GRIS_DOUX,
         align=WD_ALIGN_PARAGRAPH.CENTER, avant=0, apres=0)


def sommaire(doc, titre="Sommaire", profondeur="1-2"):
    """Champ TOC : Word le renseigne avec Ctrl+A puis F9."""
    doc.add_page_break()
    _titre(doc, titre, 1, 19, avant=0, apres=4)
    _filet_sous_titre(doc, apres=10)

    para(doc, "Le sommaire se met à jour dans Word : sélectionner tout le document "
              "avec Ctrl+A, puis appuyer sur F9 et choisir la mise à jour de toute "
              "la table.",
         taille=9, italique=True, couleur=GRIS_DOUX, apres=14)

    pt = doc.add_paragraph()
    pt.paragraph_format.space_before = Pt(0)
    pt.paragraph_format.space_after = Pt(0)
    _champ(pt, f'TOC \\o "{profondeur}" \\h \\z \\u', taille=11, couleur=TEXTE_COUL)
    return pt


def saut(doc):
    doc.add_page_break()
